from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

# Add a title
title = doc.add_heading('Document Creation Example', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a paragraph with bold and italic text
paragraph = doc.add_paragraph('This is a sample document created using the python-docx library.')
run = paragraph.runs[0]
run.bold = True
run.italic = True

# Add a heading
doc.add_heading('Section 1: Introduction', level=2)

# Add a bulleted list
list_paragraph = doc.add_paragraph()
list_paragraph.add_run('Bullet 1').bold = True
list_paragraph.add_run(' - This is the first bullet point.')
list_paragraph.add_run('\n')
list_paragraph.add_run('Bullet 2').bold = True
list_paragraph.add_run(' - This is the second bullet point.')


# Add an image
doc.add_heading('Section 3: Image', level=2)
doc.add_paragraph('Here is an image:')

# Save the document
doc.save('example_document.docx')