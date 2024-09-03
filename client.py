import tkinter as tk
from tkinter import filedialog, font, colorchooser, simpledialog, messagebox
from PIL import Image, ImageTk
import webbrowser
import os
import sys
import subprocess
from docx import Document
from docx.shared import RGBColor, Pt

class TextEditor:
    def __init__(self, root):
        self.root = root
        self.root.title("Collaborative Word")
        
        # Set the window to fullscreen
        #self.root.attributes('-topmost', True)
        self.root.state('zoomed')
        self.root.bind('<Escape>', self.toggle_fullscreen)  # Bind Escape key to exit fullscreen

        # Initialize font variables
        self.font_var = tk.StringVar(value="Arial")
        self.font_size_var = tk.StringVar(value="12")
        
        # Initialize file variable
        self.current_file = None
        self.is_file_modified = False

        # Toolbar and text area setup
        self.toolbar = tk.Frame(self.root, bd=1, relief=tk.RAISED)
        self.toolbar.pack(side=tk.TOP, fill=tk.X)
        self.text_area = tk.Text(self.root, wrap='word', undo=True)
        self.text_area.pack(fill='both', expand=True)

        # Initialize image list for holding image references
        self.text_area.image_list = []

        # Create toolbar icons and functionality
        self.add_toolbar_icons()

        # Bind events to track file modifications
        self.text_area.bind("<<Modified>>", self.on_text_modified)

        # Bind keyboard shortcuts for undo and redo
        self.root.bind('<Control-z>', self.undo)
        self.root.bind('<Control-y>', self.redo)

        # Create a label for the word counter
        self.word_count_label = tk.Label(self.toolbar, text="Words: 0")
        self.word_count_label.pack(side=tk.RIGHT, padx=10)

        # Bind events to track file modifications and word count
        self.text_area.bind("<<Modified>>", self.on_text_modified)
        root.protocol("WM_DELETE_WINDOW", self.exit_file)

    def toggle_fullscreen(self, event=None):
        """Toggle fullscreen mode."""
        is_fullscreen = self.root.attributes('-fullscreen')
        self.root.attributes('-fullscreen', not is_fullscreen)


    def add_toolbar_icons(self):
        """Adds icons to the toolbar for various functionalities."""
        icon_size = (24, 24)
        
        # Icons for toolbar buttons
        icons = {
            "file": Image.open("static/save-icon.png").resize(icon_size, Image.LANCZOS),
            "clipboard": Image.open("static/clipboard-icon.png").resize(icon_size, Image.LANCZOS),
            "bold": Image.open("static/bold-icon.png").resize(icon_size, Image.LANCZOS),
            "italic": Image.open("static/itallic-icon.png").resize(icon_size, Image.LANCZOS),
            "underline": Image.open("static/underline-icon.png").resize(icon_size, Image.LANCZOS),
            "text_color": Image.open("static/textcolor-icon.png").resize(icon_size, Image.LANCZOS),
            "insert_image": Image.open("static/image-icon.png").resize(icon_size, Image.LANCZOS),
            "insert_link": Image.open("static/link-icon.png").resize(icon_size, Image.LANCZOS),
            "change_font": Image.open("static/font-icon.png").resize(icon_size, Image.LANCZOS),
            "undo": Image.open("static/undo-icon.png").resize(icon_size, Image.LANCZOS),
            "redo": Image.open("static/redo-icon.png").resize(icon_size, Image.LANCZOS)
        }

        # Convert images to PhotoImage objects
        self.icons = {name: ImageTk.PhotoImage(icon) for name, icon in icons.items()}

        # File Button with Dropdown Menu
        file_button = tk.Menubutton(self.toolbar, image=self.icons["file"], relief=tk.RAISED)
        file_button.pack(side=tk.LEFT, padx=2, pady=2)
        file_button.menu = tk.Menu(file_button, tearoff=0)
        file_button["menu"] = file_button.menu
        self.create_file_menu(file_button.menu)

        # Clipboard Button with Dropdown Menu
        clipboard_button = tk.Menubutton(self.toolbar, image=self.icons["clipboard"], relief=tk.RAISED)
        clipboard_button.pack(side=tk.LEFT, padx=2, pady=2)
        clipboard_button.menu = tk.Menu(clipboard_button, tearoff=0)
        clipboard_button["menu"] = clipboard_button.menu
        self.create_clipboard_menu(clipboard_button.menu)

        # Font Button with Dropdown Menu
        font_button = tk.Menubutton(self.toolbar, image=self.icons["change_font"], relief=tk.RAISED)
        font_button.pack(side=tk.LEFT, padx=2, pady=2)
        font_button.menu = tk.Menu(font_button, tearoff=0)
        font_button["menu"] = font_button.menu
        self.create_font_menu(font_button.menu)

        # Other Toolbar Buttons
        tk.Button(self.toolbar, image=self.icons["bold"], command=self.make_bold).pack(side=tk.LEFT, padx=2, pady=2)
        tk.Button(self.toolbar, image=self.icons["italic"], command=self.make_italic).pack(side=tk.LEFT, padx=2, pady=2)
        tk.Button(self.toolbar, image=self.icons["underline"], command=self.make_underline).pack(side=tk.LEFT, padx=2, pady=2)
        tk.Button(self.toolbar, image=self.icons["text_color"], command=self.change_text_color).pack(side=tk.LEFT, padx=2, pady=2)
        tk.Button(self.toolbar, image=self.icons["insert_image"], command=self.insert_image).pack(side=tk.LEFT, padx=2, pady=2)
        tk.Button(self.toolbar, image=self.icons["insert_link"], command=self.insert_link).pack(side=tk.LEFT, padx=2, pady=2)
        tk.Button(self.toolbar, image=self.icons["undo"], command=self.undo).pack(side=tk.LEFT, padx=2, pady=2)
        tk.Button(self.toolbar, image=self.icons["redo"], command=self.redo).pack(side=tk.LEFT, padx=2, pady=2)

    def create_file_menu(self, menu):
        """Creates a dropdown menu for file operations."""
        menu.add_command(label="New", command=self.new_file)
        menu.add_command(label="Open", command=self.open_file)
        menu.add_command(label="Save", command=self.save_file)
        menu.add_command(label="Save As", command=self.save_as_file)
        menu.add_command(label="Exit", command=self.exit_file)

    def create_clipboard_menu(self, menu):
        """Creates a dropdown menu for clipboard operations."""
        menu.add_command(label="Copy", command=self.copy_text)
        menu.add_command(label="Cut", command=self.cut_text)
        menu.add_command(label="Paste", command=self.paste_text)

    def create_font_menu(self, menu):
        """Creates a dropdown menu for selecting font family and size."""
        font_families = sorted(font.families())

        # Create submenu for font families
        font_family_menu = tk.Menu(menu, tearoff=False)
        for family in font_families:
            font_family_menu.add_radiobutton(label=family, variable=self.font_var, command=self.apply_font)
        
        # Create submenu for font sizes
        font_size_menu = tk.Menu(menu, tearoff=False)
        for size in range(8, 72, 2):
            font_size_menu.add_radiobutton(label=str(size), variable=self.font_size_var, command=self.apply_font)
        
        menu.add_cascade(label="Font Family", menu=font_family_menu)
        menu.add_cascade(label="Font Size", menu=font_size_menu)

    def apply_font(self):
        selected_font = font.Font(family=self.font_var.get(), size=self.font_size_var.get())
        tag_name = f"font_{self.font_var.get()}_{self.font_size_var.get()}"

        self.text_area.tag_configure(tag_name, font=selected_font)

        try:
            self.text_area.tag_add(tag_name, "sel.first", "sel.last")
        except tk.TclError:
            pass  # Ignore if no text is selected

    def change_text_color(self):
        color = colorchooser.askcolor()[1]
        if color:
            tag_name = f"color_{color}"
            self.text_area.tag_configure(tag_name, foreground=color)
            try:
                self.text_area.tag_add(tag_name, "sel.first", "sel.last")
            except tk.TclError:
                pass

    def cut_text(self):
        self.copy_text()
        self.text_area.delete("sel.first", "sel.last")

    def copy_text(self):
        try:
            selected_text = self.text_area.get("sel.first", "sel.last")
            self.root.clipboard_clear()
            self.root.clipboard_append(selected_text)
        except tk.TclError:
            pass

    def paste_text(self):
        try:
            clipboard_text = self.root.clipboard_get()
            self.text_area.insert(tk.INSERT, clipboard_text)
        except tk.TclError:
            pass

    def new_file(self):
        if self.is_file_modified:
            if not self.ask_save_changes():
                return
        self.text_area.delete(1.0, tk.END)
        self.current_file = None
        self.is_file_modified = False

    def open_file(self):
        if self.is_file_modified:
            if not self.ask_save_changes():
                return
        file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
        if file_path:
            doc = Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
            self.text_area.delete(1.0, tk.END)
            self.text_area.insert(tk.END, content)
            self.current_file = file_path
            self.is_file_modified = False

    def save_to_docx_with_formatting(self, doc, content, tags_info):
        """
        Save the text content along with its formatting to a .docx file.
        """
        doc = Document()
        paragraph = doc.add_paragraph()
        # Initialize the position tracker
        current_position = 0

        # Iterate through the tags and content to apply formatting
        for tag_range, tag_dict in tags_info.items():
            print(tag_range)
            print(tag_dict)
            start_idx, end_idx = tag_range
            text_segment = content[start_idx:end_idx]

            # Add the text to the document
            run = paragraph.add_run(text_segment)
            doc.save(self.current_file)
            # Apply the formatting
            #runf = self.apply_run_formatting(run, tag_dict)

        

    def apply_run_formatting(self, run, tag_dict):
        """
        Apply the extracted formatting from tags to a document run.
        """
        if "bold" in tag_dict:
            run.bold = True
        if "italic" in tag_dict:
            run.italic = True
        if "underline" in tag_dict:
            run.underline = True
        if "foreground" in tag_dict:
            # Convert color from hex to RGB
            hex_color = tag_dict["foreground"]
            run.font.color.rgb = RGBColor(int(hex_color[1:3], 16), int(hex_color[3:5], 16), int(hex_color[5:7], 16))
        if "font" in tag_dict:
            font_family, font_size = tag_dict["font"].split("_")[1:]
            run.font.name = font_family
            run.font.size = Pt(int(font_size))

    def get_text_with_tags(self):
        """
        Get the text along with all tags and their ranges.
        """
        content = self.text_area.get("1.0", "end-1c")  # All text without the last newline
        tags_info = {}
        for tag in self.text_area.tag_names():
            ranges = self.text_area.tag_ranges(tag)
            for i in range(0, len(ranges), 2):
                start = self.text_area.index(ranges[i])
                end = self.text_area.index(ranges[i + 1])
                start_idx = self.text_area.count("1.0", start)[0]
                end_idx = self.text_area.count("1.0", end)[0]
                if (start_idx, end_idx) not in tags_info:
                    tags_info[(start_idx, end_idx)] = {}
                tags_info[(start_idx, end_idx)][tag] = self.text_area.tag_cget(tag, "foreground") if tag.startswith("color_") else True
        return content, tags_info

    def save_file(self):
        if self.current_file:
            
            content, tags_info = self.get_text_with_tags()
            self.save_to_docx_with_formatting(content, tags_info)
            print(tags_info)
            
            self.is_file_modified = False
        else:
            self.save_as_file()

    def save_as_file(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
        if file_path:
            doc = Document()
            content, tags_info = self.get_text_with_tags()
            self.save_to_docx_with_formatting(doc, content, tags_info)
            doc.save(file_path)
            self.current_file = file_path
            self.is_file_modified = False

    def exit_file(self):
        if self.is_file_modified:
            if not self.ask_save_changes():
                return
        
        # Close the current application
        self.root.destroy()
        
        # Start homepage.py
        script_path = os.path.abspath("homepage.py")
        subprocess.Popen([sys.executable, script_path])
        
        # Optionally, you could terminate the current script if needed
        # sys.exit()

    def ask_save_changes(self):
        """Prompt the user to save changes if the file is modified."""
        response = messagebox.askyesnocancel("Unsaved Changes", "You have unsaved changes. Do you want to save them?")
        if response:  # Yes
            self.save_file()
            return True
        elif response is False:  # No
            return True
        else:  # Cancel
            return False
        
        
    def update_word_count(self):
        text_content = self.text_area.get(1.0, tk.END)
        word_count = len(text_content.split())
        self.word_count_label.config(text=f"Words: {word_count}")
        
    def on_text_modified(self, event=None):
        self.is_file_modified = self.text_area.edit_modified()
        self.text_area.edit_modified(False)
        self.update_word_count()

    def undo(self, event=None):
        """Undo the last operation."""
        self.text_area.edit_undo()
        self.update_word_count()

    def redo(self, event=None):
        """Redo the last undone operation."""
        self.text_area.edit_redo()
        self.update_word_count()

    def make_bold(self):
        current_tags = self.text_area.tag_names("sel.first")
        if "italic" in current_tags and "underline" in current_tags:
            self.make_bolditalicunderline()
        elif "italic" in current_tags:
            self.make_bolditalic()
        elif "underline" in current_tags:
            self.make_boldunderline()
        else:
            bold_font = font.Font(self.text_area, self.text_area.cget("font"))
            bold_font.configure(weight="bold")
            self.text_area.tag_configure("bold", font=bold_font)

            if "bold" in current_tags:
                self.text_area.tag_remove("bold", "sel.first", "sel.last")
            else:
                self.text_area.tag_add("bold", "sel.first", "sel.last")

    def make_italic(self):
        current_tags = self.text_area.tag_names("sel.first")
        if "bold" in current_tags and "underline" in current_tags:
            self.make_bolditalicunderline()
        elif "bold" in current_tags:
            self.make_bolditalic()
        elif "underline" in current_tags:
            self.make_italicunderline()
        else:
            italic_font = font.Font(self.text_area, self.text_area.cget("font"))
            italic_font.configure(slant="italic")
            self.text_area.tag_configure("italic", font=italic_font)

            if "italic" in current_tags:
                self.text_area.tag_remove("italic", "sel.first", "sel.last")
            else:
                self.text_area.tag_add("italic", "sel.first", "sel.last")

    def make_underline(self):
        current_tags = self.text_area.tag_names("sel.first")
        if "bolditalic" in current_tags:
            self.make_bolditalicunderline()
        elif "bold" in current_tags:
            self.make_boldunderline()
        elif "italic" in current_tags:
            self.make_italicunderline()
        else:
            underline_font = font.Font(self.text_area, self.text_area.cget("font"))
            underline_font.configure(underline=True)
            self.text_area.tag_configure("underline", font=underline_font)

            if "underline" in current_tags:
                self.text_area.tag_remove("underline", "sel.first", "sel.last")
            else:
                self.text_area.tag_add("underline", "sel.first", "sel.last")

    def make_bolditalic(self):
        bolditalic_font = font.Font(self.text_area, self.text_area.cget("font"))
        bolditalic_font.configure(weight="bold", slant="italic")
        self.text_area.tag_configure("bolditalic", font=bolditalic_font)

        current_tags = self.text_area.tag_names("sel.first")
        if "bolditalic" in current_tags:
            self.text_area.tag_remove("bolditalic", "sel.first", "sel.last")
        else:
            self.text_area.tag_add("bolditalic", "sel.first", "sel.last")

        if "bold" in current_tags:
            self.text_area.tag_remove("bold", "sel.first", "sel.last")
        if "italic" in current_tags:
            self.text_area.tag_remove("italic", "sel.first", "sel.last")

    def make_boldunderline(self):
        boldunderline_font = font.Font(self.text_area, self.text_area.cget("font"))
        boldunderline_font.configure(weight="bold", underline=True)
        self.text_area.tag_configure("boldunderline", font=boldunderline_font)

        current_tags = self.text_area.tag_names("sel.first")
        if "boldunderline" in current_tags:
            self.text_area.tag_remove("boldunderline", "sel.first", "sel.last")
        else:
            self.text_area.tag_add("boldunderline", "sel.first", "sel.last")

        if "bold" in current_tags:
            self.text_area.tag_remove("bold", "sel.first", "sel.last")
        if "underline" in current_tags:
            self.text_area.tag_remove("underline", "sel.first", "sel.last")

    def make_italicunderline(self):
        italicunderline_font = font.Font(self.text_area, self.text_area.cget("font"))
        italicunderline_font.configure(slant="italic", underline=True)
        self.text_area.tag_configure("italicunderline", font=italicunderline_font)

        current_tags = self.text_area.tag_names("sel.first")
        if "italicunderline" in current_tags:
            self.text_area.tag_remove("italicunderline", "sel.first", "sel.last")
        else:
            self.text_area.tag_add("italicunderline", "sel.first", "sel.last")

        if "italic" in current_tags:
            self.text_area.tag_remove("italic", "sel.first", "sel.last")
        if "underline" in current_tags:
            self.text_area.tag_remove("underline", "sel.first", "sel.last")

    def make_bolditalicunderline(self):
        bolditalicunderline_font = font.Font(self.text_area, self.text_area.cget("font"))
        bolditalicunderline_font.configure(weight="bold", slant="italic", underline=True)
        self.text_area.tag_configure("bolditalicunderline", font=bolditalicunderline_font)
        current_tags = self.text_area.tag_names("sel.first")
        if "bolditalicunderline" in current_tags:
            self.text_area.tag_remove("bolditalicunderline", "sel.first", "sel.last")
        else:
            self.text_area.tag_add("bolditalicunderline", "sel.first", "sel.last")

        if "bold" in current_tags:
            self.text_area.tag_remove("bold", "sel.first", "sel.last")
        if "italic" in current_tags:
            self.text_area.tag_remove("italic", "sel.first", "sel.last")
        if "underline" in current_tags:
            self.text_area.tag_remove("underline", "sel.first", "sel.last")

    def insert_hyperlink_with_hover(self):
        url = simpledialog.askstring("Insert Hyperlink", "Enter URL:")
        if url:
            display_text = simpledialog.askstring("Insert Hyperlink", "Enter the display text:")
            if not display_text:
                display_text = url  

        link = self.hyperlink(self.root, url, text=display_text)
        self.text_area.window_create(tk.INSERT, window=link)

    def hyperlink(self, root, link: str, **kwargs) -> tk.Label:
        for i in (("fg", "magenta"), ("text", "Hyperlink!"), ("font", "None 10"), ("cursor", "hand2")):
            kwargs.setdefault(i[0], i[1])

        label = tk.Label(root, **kwargs)
        label.link = link

        label.bind("<Button-1>", lambda e: webbrowser.open(e.widget.link))
        label.bind("<Enter>", lambda e: e.widget.configure(font=e.widget.cget("font") + " underline"))
        label.bind("<Leave>", lambda e: e.widget.configure(font=e.widget.cget("font")[:-10]))

        return label 

    def insert_image(self):
        file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png;*.jpg;*.jpeg;*.gif")])
        if file_path:
            image = Image.open(file_path)
            image.thumbnail((100, 100))
            photo = ImageTk.PhotoImage(image)
            self.text_area.image_create(tk.END, image=photo)
            self.text_area.image_list.append(photo)  # Save reference

    def insert_link(self):
        """Insert a hyperlink into the text area."""
        url = simpledialog.askstring("Insert Link", "Enter URL:")
        if url:
            self.text_area.insert(tk.INSERT, url)
            self.text_area.tag_add("hyperlink", tk.INSERT + "-%dc" % len(url), tk.INSERT)
            self.text_area.tag_config("hyperlink", foreground="blue", underline=True)
            self.text_area.tag_bind("hyperlink", "<Button-1>", lambda e: webbrowser.open(url))

if __name__ == "__main__":
    root = tk.Tk()
    app = TextEditor(root)
    root.mainloop()
